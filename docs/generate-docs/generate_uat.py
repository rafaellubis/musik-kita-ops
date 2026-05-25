r"""
Generate UAT-Musik-KITA-v1.0.docx
Jalankan: C:\laragon\bin\python\python-3.13\python.exe docs/generate-docs/generate_uat.py
Output  : UAT-Musik-KITA-v1.0.docx (di root project)
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import date
import os

OUTPUT_PATH = os.path.join(os.path.dirname(__file__), '..', '..', 'UAT-Musik-KITA-v1.0.docx')

# ─── Helpers ──────────────────────────────────────────────────────────────────

def set_font(run, size=10, bold=False, color=None):
    run.font.name = 'Calibri'
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = 'Calibri'
        run.font.size = Pt(14 if level == 1 else 12)
        run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    return p

def set_col_width(table, col_widths_cm):
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            if i < len(col_widths_cm):
                cell.width = Cm(col_widths_cm[i])

def shade_header_row(table):
    """Warna biru muda untuk baris header tabel."""
    tr = table.rows[0]._tr
    trPr = tr.get_or_add_trPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'BDD7EE')
    trPr.append(shd)

def add_scenario_table(doc, module_name, scenarios):
    """
    Tambah satu tabel modul.
    scenarios: list of tuples (no, skenario, langkah_uji, expected_result)
    """
    add_heading(doc, module_name, level=2)
    headers = ['No', 'Skenario', 'Langkah Uji', 'Expected Result', 'Status (v/x)', 'Catatan']
    widths  = [1.0, 3.5, 5.5, 4.5, 1.5, 2.0]  # total ~18cm

    table = doc.add_table(rows=1, cols=6)
    table.style = 'Table Grid'
    shade_header_row(table)

    # Header row
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        p = hdr[i].paragraphs[0]
        run = p.add_run(h)
        set_font(run, size=9, bold=True)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Data rows
    for sc in scenarios:
        row = table.add_row()
        values = [str(sc[0]), sc[1], sc[2], sc[3], '', '']
        for i, val in enumerate(values):
            p = row.cells[i].paragraphs[0]
            run = p.add_run(val)
            set_font(run, size=9)
            if i == 0:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    set_col_width(table, widths)
    doc.add_paragraph()  # spasi antar tabel


# ─── Data Skenario ─────────────────────────────────────────────────────────────

M01_SCENARIOS = [
    (1,  'Tambah instrumen baru',
         '1. Master Data > Instrumen > Tambah\n2. Isi nama instrumen, simpan',
         'Instrumen muncul di daftar, bisa dipilih saat buat paket'),
    (2,  'Edit instrumen',
         '1. Pilih instrumen > Edit\n2. Ubah nama > Simpan',
         'Perubahan tersimpan dan tampil di daftar'),
    (3,  'Tambah paket baru (Reguler)',
         '1. Master Data > Paket > Tambah\n2. Isi kode, instrumen, class_type=REGULER, grade, durasi, harga\n3. Simpan',
         'Paket muncul di daftar dengan harga benar'),
    (4,  'Tambah paket baru (Hobby 45 menit)',
         '1. Buat paket class_type=HOBBY, durasi 45 menit, harga 450.000\n2. Simpan',
         'Paket tersimpan dengan durasi dan harga sesuai'),
    (5,  'Edit harga paket (Owner only)',
         '1. Login sebagai Owner > Paket > Edit > Ubah harga\n2. Simpan',
         'Harga berubah. Admin tidak bisa akses menu edit harga'),
    (6,  'Tambah data guru baru',
         '1. Master Data > Guru > Tambah\n2. Isi nama, kode, instrumen yang diajar, info bank\n3. Simpan',
         'Guru muncul di daftar, bisa dipilih saat daftar murid'),
    (7,  'Edit data guru (info bank)',
         '1. Pilih guru > Edit\n2. Ubah nama bank / nomor rekening > Simpan',
         'Info bank tersimpan, tampil di header slip honor cetak'),
    (8,  'Nonaktifkan guru (ada sesi historis)',
         '1. Guru dengan sesi historis > Nonaktifkan\n2. Konfirmasi',
         'Guru berubah status tidak aktif. Tidak bisa dihapus jika ada sesi historis'),
    (9,  'Tambah ruangan baru',
         '1. Master Data > Ruangan > Tambah\n2. Isi kode, nama, kapasitas, instrumen yang didukung\n3. Simpan',
         'Ruangan muncul di daftar dan bisa dipilih saat atur jadwal'),
    (10, 'Tambah hari libur nasional (tanpa pengganti)',
         '1. Master Data > Hari Libur > Tambah\n2. Tipe: Nasional, tanpa replacement_date\n3. Simpan',
         'Hari libur tersimpan. Sesi yang jatuh pada tanggal ini berstatus LIBUR'),
    (11, 'Tambah hari libur nasional (dengan tanggal pengganti)',
         '1. Tambah hari libur\n2. Isi replacement_date (minggu ke-5 bulan yang sama)\n3. Simpan',
         'Saat generate sesi, ada sesi LIBUR + sesi pengganti di tanggal replacement'),
    (12, 'Tambah hari libur Internal (Konser KITA)',
         '1. Tipe: Internal, is_honor_paid=false\n2. replacement_date dibiarkan kosong\n3. Simpan',
         'Hari libur tersimpan. Honor guru Rp 0 untuk sesi ini. Field pengganti tidak bisa diisi'),
]

M02_SCENARIOS = [
    (13, 'Daftar murid baru (status Calon)',
         '1. Murid > Tambah Murid\n2. Isi semua data wajib (nama, gender, orang tua, dll)\n3. Simpan',
         'Murid tersimpan dengan status Calon, kode M-YYYY-NNNN terotomatis'),
    (14, 'Jadwalkan sesi trial',
         '1. Detail murid (Calon) > Jadwalkan Trial\n2. Isi tanggal, jam, guru, paket diminati\n3. Simpan',
         'Status murid berubah ke Trial. Sesi trial muncul di jadwal hari tsb'),
    (15, 'Input absensi trial -- murid HADIR',
         '1. Absensi > cari sesi trial > set HADIR\n2. Simpan',
         'Honor guru terbayar penuh sesuai paket. Sesi berstatus HADIR'),
    (16, 'Input absensi trial -- murid NO-SHOW',
         '1. Set sesi trial > HANGUS (tidak hadir tanpa konfirmasi)\n2. Simpan',
         'Honor guru Rp 0 (TRIAL_NS). Sesi berstatus HANGUS'),
    (17, 'Konversi Trial -> Aktif',
         '1. Detail murid (Trial) > Konversi ke Aktif\n2. Isi paket, guru, ruang, jadwal mingguan\n3. Konfirmasi',
         'Status murid Aktif. Invoice registrasi Rp 250.000 + SPP bulan berjalan ter-generate otomatis'),
    (18, 'Skip trial -- langsung Aktif (walk_in)',
         '1. Murid (Calon) > Skip Trial > pilih alasan: walk_in\n2. Isi paket, guru, jadwal\n3. Konfirmasi',
         'Status langsung Aktif. Reason tersimpan di histori status. Invoice registrasi + SPP ter-generate'),
    (19, 'Murid tidak melanjutkan setelah trial',
         '1. Detail murid (Trial) > Mundurkan\n2. Isi alasan > Konfirmasi',
         'Status murid berubah ke Mengundurkan Diri. Histori tercatat'),
]

M03_SCENARIOS = [
    (20, 'Generate sesi bulanan -- bulan normal (tidak ada libur)',
         '1. Penjadwalan > Generate Sesi > pilih bulan\n2. Konfirmasi generate',
         'Setiap murid Aktif mendapat 4 sesi sesuai jadwal mingguan tetap'),
    (21, 'Generate sesi bulanan -- ada hari libur nasional tanpa pengganti',
         '1. Pastikan ada holiday tanpa replacement_date di bulan tsb\n2. Generate sesi',
         'Murid yang jadwal hariannya libur mendapat sesi LIBUR (tidak diganti). Total bisa 3 sesi bulan itu'),
    (22, 'Generate sesi bulanan -- ada hari libur nasional dengan pengganti',
         '1. Pastikan ada holiday dengan replacement_date di bulan tsb\n2. Generate sesi',
         'Ada sesi LIBUR + sesi pengganti di tanggal replacement. Total tetap 4 sesi'),
    (23, 'Generate sesi -- minggu ke-5 tidak ditambah secara reguler',
         '1. Generate sesi bulan yang punya 5 minggu hari tersebut\n2. Cek hasil',
         'Maksimal 4 sesi per murid. Minggu ke-5 tidak ter-generate kecuali sebagai sesi pengganti'),
    (24, 'Cek sesi murid cuti tidak ter-generate',
         '1. Ajukan cuti untuk murid Aktif\n2. Generate sesi bulan cuti',
         'Murid yang sedang cuti tidak mendapat sesi di bulan tsb'),
    (25, 'Pindah jadwal mingguan tetap (dalam bulan berjalan)',
         '1. Detail murid > Jadwal > Edit jadwal\n2. Ubah hari/jam (dalam bulan yang sama)\n3. Simpan',
         'Jadwal baru berlaku. Sesi bulan berjalan menyesuaikan'),
]

M04_SCENARIOS = [
    (26, 'Input absensi HADIR',
         '1. Absensi > cari sesi hari ini > set HADIR\n2. Simpan',
         'Sesi berstatus HADIR. Honor terhitung otomatis (H_REG)'),
    (27, 'Input absensi HADIR_TERLAMBAT',
         '1. Set sesi > HADIR_TERLAMBAT, isi menit terlambat\n2. Simpan',
         'Sesi berstatus HADIR_TERLAMBAT. Honor tetap penuh (H_REG)'),
    (28, 'Input absensi IZIN_RESCHEDULE (izin pertama bulan ini, info >=5 jam)',
         '1. Set sesi > IZIN_RESCHEDULE\n2. Konfirmasi info >=5 jam + izin pertama\n3. Simpan',
         'Sesi berstatus IZIN_RESCHEDULE. Tombol buat sesi pengganti aktif. Honor sesi ini Rp 0 (H_IZIN)'),
    (29, 'Buat sesi pengganti (reschedule) via mini-modal',
         '1. Klik "Buat Sesi Pengganti" pada sesi IZIN_RESCHEDULE\n2. Isi tanggal, jam, ruang pengganti\n3. Konfirmasi',
         'Sesi pengganti terbuat. Konflik guru+ruang dicek otomatis. Honor dibayar di sesi pengganti'),
    (30, 'Input absensi HANGUS (no-show tanpa info)',
         '1. Set sesi > HANGUS\n2. Simpan',
         'Sesi berstatus HANGUS. Honor guru tetap dibayar (H_HANGUS)'),
    (31, 'Input absensi IZIN_RESCHEDULE ke-2 (bulan yang sama)',
         '1. Murid sudah punya 1 izin di bulan ini\n2. Set sesi ke-2 > sistem tidak izinkan reschedule\n3. Sistem saran IZIN_VIDEO',
         'Izin ke-2 tidak bisa reschedule. Diarahkan ke IZIN_VIDEO'),
    (32, 'Input absensi DIGANTI -- guru pengganti',
         '1. Set sesi > DIGANTI\n2. Pilih guru pengganti dari dropdown\n3. Simpan',
         'Honor dialihkan ke guru pengganti (H_PENG). Guru utama honor Rp 0'),
    (33, 'Sesi LIBUR -- honor guru',
         '1. Cek sesi yang berstatus LIBUR (hari libur nasional)\n2. Lihat honor_amount di detail sesi',
         'Honor guru Rp penuh (H_LIBUR). Bukan Rp 0'),
    (34, 'Absensi Kids Class -- input per murid dalam grup',
         '1. Absensi > cari sesi Kids Class\n2. Input status per murid (bisa beda-beda)\n3. Simpan',
         'Setiap murid punya status absensi sendiri. Honor guru = jumlah murid hadir x Rp 42.500'),
]

M05_SCENARIOS = [
    (35, 'Generate SPP bulanan (tanggal 1)',
         '1. Keuangan > Generate SPP > pilih bulan/tahun > Konfirmasi',
         'Invoice SPP ter-generate untuk semua murid Aktif dengan primary enrollment. Invoice duplikat tidak dibuat'),
    (36, 'Generate SPP -- skip murid Kids Class Bundle',
         '1. Generate SPP bulan berjalan\n2. Cek murid Kids Class Bundle di list',
         'Murid Kids Class Bundle tidak mendapat invoice SPP bulanan (sudah tercover cicilan)'),
    (37, 'Catat pembayaran SPP (CASH)',
         '1. Detail invoice > Catat Pembayaran\n2. Pilih metode CASH, isi nominal, tanggal\n3. Simpan',
         'Pembayaran tercatat. Kuitansi KW/YYYY/MM/NNNN ter-generate. Status invoice update'),
    (38, 'Catat pembayaran (TRANSFER)',
         '1. Catat pembayaran metode TRANSFER\n2. Upload foto bukti transfer\n3. Simpan',
         'Pembayaran tersimpan dengan bukti foto. Status invoice update'),
    (39, 'Catat pembayaran (QRIS)',
         '1. Catat pembayaran metode QRIS\n2. Simpan',
         'QRIS tercatat sebagai metode. Status invoice update'),
    (40, 'Apply denda keterlambatan (tanggal 11+)',
         '1. Keuangan > Apply Denda > pilih bulan (invoice unpaid sudah lewat tanggal 10)\n2. Konfirmasi',
         'Item DENDA Rp 5.000 x hari terlambat ditambah ke invoice. Total tagihan bertambah'),
    (41, 'Cetak invoice / kuitansi (format A4)',
         '1. Detail invoice > Cetak\n2. Browser membuka halaman print\n3. Ctrl+P > simpan PDF',
         'Halaman cetak A4 tanpa navigasi. Tombol cetak muncul, tersembunyi saat print'),
    (42, 'Tambah item manual ke invoice (dari katalog)',
         '1. Detail invoice (UNPAID) > Tambah Item\n2. Pilih item dari dropdown katalog (misal CUTI)\n3. Simpan',
         'Item baru ditambahkan ke invoice. Total tagihan bertambah'),
    (43, 'Tambah diskon NOMINAL ke item invoice',
         '1. Detail invoice > pilih item yang akan didiskon > Tambah Diskon\n2. Tipe NOMINAL, isi nominal, wajib isi alasan\n3. Simpan',
         'Item DISKON terbuat sebagai child item. Total invoice berkurang sesuai nominal'),
    (44, 'Tambah diskon PERCENT ke item invoice',
         '1. Tipe PERCENT > isi persentase, wajib isi alasan > Simpan',
         'Diskon dihitung dari amount item parent. Total invoice berkurang'),
    (45, 'Void pembayaran (Owner only)',
         '1. Login Owner > Detail invoice > klik void pada payment\n2. Isi alasan void > Konfirmasi',
         'Pembayaran di-void. Status invoice kembali UNPAID/PARTIAL. Admin tidak bisa void'),
    (46, 'Generate 3 invoice cicilan Kids Class Bundle',
         '1. Detail murid (Aktif, KIDS_CLASS_BUNDLE) > Generate Cicilan Bundle\n2. Isi tanggal mulai program\n3. Konfirmasi',
         '3 invoice cicilan terbuat: bulan 1, 2, 4. Total 3 termin = Rp 2.180.000. Tidak bisa generate ulang'),
    (47, 'Cek status tagihan outstanding',
         '1. Keuangan > Daftar Invoice > filter status UNPAID\n2. Lihat daftar dan aging',
         'Invoice UNPAID tampil dengan info overdue. Murid tunggakan >1 bulan muncul warning'),
]

M06_SCENARIOS = [
    (48, 'Kalkulasi honor guru bulanan (H-2 sebelum akhir bulan)',
         '1. Honor > Kalkulasi Honor > pilih bulan/tahun\n2. Konfirmasi',
         'Honor terhitung untuk semua guru dari sesi yang sudah ada status final. Slip DRAFT terbuat'),
    (49, 'Lihat rincian slip honor guru',
         '1. Honor > pilih guru > Detail Slip\n2. Lihat per sesi dan per murid',
         'Rincian sesi tampil: tanggal, murid, honor_code, honor_amount. Total akurat'),
    (50, 'Edit komponen honor manual (transport, lain-lain)',
         '1. Slip honor (DRAFT/CALCULATED) > Edit\n2. Isi honor transport, honor lain-lain + keterangan\n3. Simpan',
         'Nilai tersimpan. Total honor = base + event + transport + lain-lain'),
    (51, 'Cetak slip honor guru',
         '1. Detail slip honor > Cetak\n2. Halaman A4 print terbuka',
         'Slip cetak berisi nama guru, periode, info bank, rincian sesi, dan total honor'),
    (52, 'Tandai honor dibayar (Owner only)',
         '1. Slip honor (CALCULATED) > Tandai Dibayar\n2. Konfirmasi',
         'Status slip berubah ke PAID. Slip terkunci dari edit. Admin tidak bisa tandai dibayar'),
    (53, 'Slip honor Kids Class -- honor flat per murid',
         '1. Lihat slip honor guru ICA (Kids Class)\n2. Cek honor per sesi',
         'Honor per sesi = jumlah murid Kids Class x Rp 42.500. Honor_code = H_KIDS'),
]

M07_SCENARIOS = [
    (54, 'Catat pengeluaran baru',
         '1. Pengeluaran > Tambah\n2. Isi kategori (Sewa/Listrik/dll), nominal, tanggal\n3. Simpan',
         'Pengeluaran tersimpan. Muncul di rekap pengeluaran bulan tsb'),
    (55, 'Lihat rekap pengeluaran bulanan',
         '1. Pengeluaran > filter bulan\n2. Lihat total per kategori',
         'Rekap menampilkan total per kategori dan grand total bulan tsb'),
]

M08_SCENARIOS = [
    (56, 'Buat event Mini Concert baru',
         '1. Event > Tambah Event\n2. Isi nama, tanggal, tipe (Mini Concert)\n3. Simpan',
         'Event terbuat dengan status DRAFT. Siap untuk pendaftaran peserta'),
    (57, 'Daftar murid ke event (Ujian + Tampil)',
         '1. Detail event > Daftar Peserta > Tambah\n2. Pilih murid, tipe: Ujian + Tampil (Rp 395.000)\n3. Simpan',
         'Murid terdaftar. Invoice event ter-generate dengan nominal Rp 395.000'),
    (58, 'Daftar murid ke event (Tampil saja)',
         '1. Pilih murid, tipe: Tampil saja (Rp 295.000)\n2. Simpan',
         'Murid terdaftar. Invoice Rp 295.000 ter-generate'),
    (59, 'Assign guru pendamping Konser KITA',
         '1. Detail event (DRAFT) > Peserta > pilih murid > Assign Guru Pendamping\n2. Pilih guru dari dropdown\n3. Simpan',
         'Guru pendamping tercatat di tabel event_participants. Bisa diubah selama event DRAFT'),
    (60, 'Input hasil ujian dan naik grade',
         '1. Event > peserta yang ikut ujian > Input Hasil\n2. Set hasil: Lulus, grade before, grade after\n3. Simpan',
         'Jika Lulus, grade murid di enrollment naik otomatis ke grade berikutnya'),
]

IMPORT_SCENARIOS = [
    (61, 'Upload file Excel template (dry-run preview)',
         '1. Import > Upload File Excel > pilih file template\n2. Submit tanpa konfirmasi (dry-run)',
         'Tabel preview tampil: data valid (hijau) dan error (merah). Belum ada data tersimpan'),
    (62, 'Konfirmasi import setelah preview bersih',
         '1. Setelah dry-run tidak ada error > klik Konfirmasi Import\n2. Tunggu proses',
         'Data murid, enrollment, dan jadwal tersimpan ke database. Laporan hasil tampil'),
    (63, 'Upload file Excel dengan baris error (nama kosong)',
         '1. Upload file dengan baris yang nama muridnya kosong\n2. Dry-run',
         'Baris error ditandai merah. Pesan error jelas. Baris valid tetap hijau'),
    (64, 'Upload file Excel dengan guru tidak dikenal',
         '1. Upload file dengan kode guru tidak ada di master data\n2. Dry-run',
         'Error: "Guru [kode] tidak ditemukan". Baris tersebut tidak diimport saat konfirmasi'),
    (65, 'Import murid Kids Class Bundle',
         '1. Upload file Excel dengan murid Kids Class Bundle\n2. Dry-run > preview\n3. Konfirmasi',
         'Murid Kids Class Bundle terimport. Jadwal mereka terbuat (tidak konflik meski guru sama)'),
    (66, 'Import murid duplikat (student_code sudah ada)',
         '1. Import ulang file yang muridnya sudah ada di database\n2. Dry-run',
         'Baris duplikat ditandai sebagai skip/warning, tidak error fatal. Murid tidak terduplikasi'),
]

KIDS_BUNDLE_SCENARIOS = [
    (67, 'Generate cicilan Kids Bundle -- first time',
         '1. Detail murid (Aktif, KIDS_CLASS_BUNDLE) > Generate Cicilan Bundle\n2. Isi tanggal mulai program (YYYY-MM-DD)\n3. Konfirmasi',
         '3 invoice cicilan terbuat: termin 1 (bulan 1), termin 2 (bulan 2), termin 3 (bulan 4). Nominal total = harga paket'),
    (68, 'Generate cicilan Kids Bundle -- sudah pernah dibuat',
         '1. Murid yang sudah punya cicilan > coba Generate Bundle lagi\n2. Konfirmasi',
         'Sistem menolak. Pesan: "Invoice cicilan sudah pernah dibuat untuk murid ini"'),
    (69, 'Catat pembayaran cicilan termin 1',
         '1. Detail invoice cicilan termin 1 > Catat Pembayaran\n2. Nominal sesuai tagihan > Simpan',
         'Termin 1 berstatus PAID. Panel progress cicilan menampilkan 1 dari 3 lunas'),
    (70, 'Lihat panel progress cicilan di detail invoice',
         '1. Buka salah satu invoice cicilan > lihat panel Progress Cicilan\n2. Lihat 3 termin',
         'Panel menampilkan 3 termin dengan status masing-masing (UNPAID/PARTIAL/PAID) dan due date'),
]


# ─── Main ──────────────────────────────────────────────────────────────────────

def build_document():
    doc = Document()

    # Margin halaman
    for section in doc.sections:
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.0)
        section.right_margin  = Cm(2.0)

    # Cover
    doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run('UAT Checklist\nSistem Administrasi Musik KITA')
    set_font(r, size=18, bold=True, color=(0x1F, 0x49, 0x7D))

    doc.add_paragraph()
    meta_items = [
        ('Versi',         'v1.0'),
        ('Tanggal',       date.today().strftime('%d %B %Y')),
        ('Tester',        'Owner'),
        ('Lingkup',       'M01 - M08, Import Excel, Kids Class Bundle'),
        ('Dikecualikan',  'Kids Class Bundle & Monthly -- dikelola manual sementara'),
    ]
    for label, value in meta_items:
        p = doc.add_paragraph()
        r1 = p.add_run(f'{label:<16}: ')
        set_font(r1, size=11, bold=True)
        r2 = p.add_run(value)
        set_font(r2, size=11)

    doc.add_paragraph()
    note = doc.add_paragraph()
    r = note.add_run(
        'PETUNJUK: Isi kolom Status dengan v (lulus) atau x (gagal) setelah '
        'menjalankan setiap skenario. Kolom Catatan diisi jika ada temuan penting.'
    )
    set_font(r, size=10)
    note.paragraph_format.space_after = Pt(12)

    doc.add_page_break()

    # Tabel per modul
    modules = [
        ('M01 - Master Data',                  M01_SCENARIOS),
        ('M02 - Pendaftaran & Trial',           M02_SCENARIOS),
        ('M03 - Penjadwalan',                   M03_SCENARIOS),
        ('M04 - Absensi',                       M04_SCENARIOS),
        ('M05 - Keuangan',                      M05_SCENARIOS),
        ('M06 - Honor Guru',                    M06_SCENARIOS),
        ('M07 - Pengeluaran',                   M07_SCENARIOS),
        ('M08 - Event',                         M08_SCENARIOS),
        ('Import Excel',                        IMPORT_SCENARIOS),
        ('KIDS Bundle - Cicilan',               KIDS_BUNDLE_SCENARIOS),
    ]

    for module_name, scenarios in modules:
        add_scenario_table(doc, module_name, scenarios)

    # Footer: tanda tangan
    doc.add_page_break()
    add_heading(doc, 'Tanda Tangan Penguji', level=2)
    doc.add_paragraph()
    sign_table = doc.add_table(rows=3, cols=2)
    sign_table.style = 'Table Grid'
    labels = [
        ['Penguji / Tester', 'Tanggal UAT'],
        ['', ''],
        ['(                                    )', ''],
    ]
    for r_idx, row_vals in enumerate(labels):
        for c_idx, val in enumerate(row_vals):
            p = sign_table.rows[r_idx].cells[c_idx].paragraphs[0]
            run = p.add_run(val)
            set_font(run, size=10, bold=(r_idx == 0))

    return doc


if __name__ == '__main__':
    doc = build_document()
    out = os.path.abspath(OUTPUT_PATH)
    doc.save(out)
    all_scenarios = [
        M01_SCENARIOS, M02_SCENARIOS, M03_SCENARIOS,
        M04_SCENARIOS, M05_SCENARIOS, M06_SCENARIOS,
        M07_SCENARIOS, M08_SCENARIOS, IMPORT_SCENARIOS,
        KIDS_BUNDLE_SCENARIOS,
    ]
    total = sum(len(s) for s in all_scenarios)
    print(f'UAT Checklist berhasil dibuat: {out}')
    print(f'   Total skenario: {total}')
