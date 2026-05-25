r"""
Generate UserManual-Admin-Musik-KITA-v1.0.docx
Jalankan: C:\laragon\bin\python\python-3.13\python.exe docs/generate-docs/generate_manual.py
Output  : UserManual-Admin-Musik-KITA-v1.0.docx (di root project)
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import date
import os

OUTPUT_PATH = os.path.join(os.path.dirname(__file__), '..', '..', 'UserManual-Admin-Musik-KITA-v1.0.docx')

# ─── Helpers ──────────────────────────────────────────────────────────────────

def set_font(run, size=10.5, bold=False, color=None, italic=False):
    run.font.name = 'Calibri'
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    sizes = {1: 16, 2: 13, 3: 11}
    for run in p.runs:
        run.font.name = 'Calibri'
        run.font.size = Pt(sizes.get(level, 11))
        run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    return p

def add_body(doc, text, indent=False):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(text)
    set_font(run)
    return p

def add_step(doc, number, text):
    p = doc.add_paragraph(style='List Number')
    p.paragraph_format.left_indent = Cm(1.0)
    run = p.add_run(text)
    set_font(run)

def add_note(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    r1 = p.add_run('Catatan: ')
    set_font(r1, bold=True, color=(0x00, 0x70, 0xC0))
    r2 = p.add_run(text)
    set_font(r2, italic=True, color=(0x00, 0x70, 0xC0))

def add_warning(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    r1 = p.add_run('Perhatian: ')
    set_font(r1, bold=True, color=(0xC0, 0x00, 0x00))
    r2 = p.add_run(text)
    set_font(r2, color=(0xC0, 0x00, 0x00))

def shade_header_row(table):
    tr = table.rows[0]._tr
    trPr = tr.get_or_add_trPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'BDD7EE')
    trPr.append(shd)

def add_simple_table(doc, headers, rows, col_widths_cm):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    shade_header_row(table)
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        p = hdr[i].paragraphs[0]
        run = p.add_run(h)
        set_font(run, size=9, bold=True)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for row_data in rows:
        row = table.add_row()
        for i, val in enumerate(row_data):
            p = row.cells[i].paragraphs[0]
            run = p.add_run(val)
            set_font(run, size=9)
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            if i < len(col_widths_cm):
                cell.width = Cm(col_widths_cm[i])
    doc.add_paragraph()


# ─── Konten Bab ───────────────────────────────────────────────────────────────

def bab1_memulai(doc):
    add_heading(doc, 'Bab 1 - Memulai', level=1)

    add_heading(doc, '1.1  Login ke Sistem', level=2)
    add_body(doc, 'Sistem Musik KITA diakses melalui browser di komputer studio (terhubung ke WiFi studio).')
    add_step(doc, 1, 'Buka browser (Chrome/Firefox) dan masukkan alamat IP studio di address bar (contoh: 192.168.1.10).')
    add_step(doc, 2, 'Tampil halaman login. Masukkan Email dan Password yang diberikan owner.')
    add_step(doc, 3, 'Klik tombol "Masuk". Anda akan langsung masuk ke halaman Dashboard.')
    add_warning(doc, 'Ganti password default segera setelah pertama kali login. Hubungi owner untuk reset jika lupa password.')

    add_heading(doc, '1.2  Mengenal Tampilan Utama', level=2)
    add_body(doc, 'Setelah login, layar terbagi menjadi tiga bagian:')
    add_step(doc, 1, 'Sidebar (kiri): menu navigasi ke semua modul -- Murid, Absensi, Keuangan, Honor Guru, dll.')
    add_step(doc, 2, 'Topbar (atas): nama pengguna yang sedang login, tombol ganti tema, dan tombol Keluar.')
    add_step(doc, 3, 'Area Konten (tengah-kanan): menampilkan data dan form sesuai menu yang dipilih.')
    add_note(doc, 'Gunakan tombol tema di pojok kanan atas untuk mengganti tampilan terang/gelap sesuai kenyamanan.')

    add_heading(doc, '1.3  Perbedaan Role Owner vs Admin', level=2)
    add_body(doc, 'Sistem memiliki tiga level akses. Sebagai Admin, Anda memiliki akses operasional harian.')
    headers = ['Fitur', 'Owner', 'Admin']
    rows = [
        ('Input absensi harian', 'Ya', 'Ya'),
        ('Catat pembayaran', 'Ya', 'Ya'),
        ('Daftar murid baru', 'Ya', 'Ya'),
        ('Generate SPP & denda', 'Ya', 'Ya'),
        ('Ubah harga paket', 'Ya', 'Tidak'),
        ('Void pembayaran', 'Ya', 'Tidak'),
        ('Tandai honor dibayar', 'Ya', 'Tidak'),
        ('Hapus master data', 'Ya', 'Tidak'),
        ('Lihat audit log', 'Ya', 'Tidak'),
    ]
    add_simple_table(doc, headers, rows, [6.0, 2.5, 2.5])


def bab2_harian(doc):
    add_heading(doc, 'Bab 2 - Tugas Harian', level=1)

    add_heading(doc, '2.1  Input Absensi Sesi Hari Ini', level=2)
    add_body(doc, 'Lakukan setelah setiap sesi selesai, atau di akhir hari.')
    add_step(doc, 1, 'Klik menu "Absensi" di sidebar.')
    add_step(doc, 2, 'Pastikan tanggal yang tampil adalah hari ini. Klik tanggal lain jika perlu.')
    add_step(doc, 3, 'Cari sesi yang ingin diisi (tampil berurutan per jam).')
    add_step(doc, 4, 'Klik tombol "Input" atau ikon status di kolom Status.')
    add_step(doc, 5, 'Pilih salah satu status: HADIR, HADIR_TERLAMBAT, IZIN_RESCHEDULE, HANGUS, DIGANTI.')
    add_step(doc, 6, 'Jika HADIR_TERLAMBAT: isi berapa menit terlambat.')
    add_step(doc, 7, 'Jika DIGANTI: pilih guru pengganti dari dropdown.')
    add_step(doc, 8, 'Klik "Simpan".')
    add_note(doc, 'Untuk IZIN_RESCHEDULE: setelah simpan, tombol "Buat Sesi Pengganti" akan aktif. Isi tanggal, jam, dan ruang pengganti lalu simpan.')

    add_heading(doc, '2.2  Catat Pembayaran Murid', level=2)
    add_step(doc, 1, 'Klik menu "Keuangan" > "Daftar Invoice".')
    add_step(doc, 2, 'Cari invoice murid yang akan bayar (gunakan kolom pencarian nama atau nomor invoice).')
    add_step(doc, 3, 'Klik nomor invoice untuk membuka detail.')
    add_step(doc, 4, 'Klik tombol "Catat Pembayaran".')
    add_step(doc, 5, 'Isi: Nominal, Metode (CASH / TRANSFER / QRIS / DEBIT), Tanggal.')
    add_step(doc, 6, 'Jika TRANSFER: upload foto bukti transfer (opsional tapi dianjurkan).')
    add_step(doc, 7, 'Klik "Simpan". Kuitansi KW/YYYY/MM/NNNN otomatis terbuat.')
    add_note(doc, 'Status invoice otomatis berubah: UNPAID ke PARTIAL (bayar sebagian) ke PAID (lunas). Kuitansi bisa dicetak dari tombol "Cetak" di detail invoice.')

    add_heading(doc, '2.3  Cek Status Tagihan Outstanding', level=2)
    add_step(doc, 1, 'Klik menu "Keuangan" > "Daftar Invoice".')
    add_step(doc, 2, 'Klik filter "Status" > pilih "UNPAID".')
    add_step(doc, 3, 'Lihat kolom "Jatuh Tempo" -- invoice dengan tanggal merah sudah lewat jatuh tempo.')
    add_step(doc, 4, 'Hubungi murid/orang tua untuk mengingatkan pembayaran.')
    add_warning(doc, 'Murid dengan tunggakan >1 bulan akan muncul peringatan di Dashboard. Segera laporkan ke owner jika ada tunggakan lama.')


def bab3_awal_bulan(doc):
    add_heading(doc, 'Bab 3 - Tugas Awal Bulan', level=1)
    add_body(doc, 'Lakukan tugas-tugas ini di awal setiap bulan.')

    add_heading(doc, '3.1  Generate SPP (Lakukan Tanggal 1)', level=2)
    add_body(doc, 'SPP = Surat Pembayaran Privat -- tagihan bulanan untuk semua murid aktif.')
    add_step(doc, 1, 'Klik menu "Keuangan" > "Daftar Invoice".')
    add_step(doc, 2, 'Klik tombol "Generate SPP".')
    add_step(doc, 3, 'Pilih Tahun dan Bulan yang akan di-generate.')
    add_step(doc, 4, 'Klik "Generate". Sistem akan membuat invoice untuk semua murid Aktif.')
    add_step(doc, 5, 'Lihat pesan hasil: berapa invoice baru dibuat, berapa skip (sudah ada).')
    add_note(doc, 'Generate SPP bersifat idempotent -- aman dijalankan ulang. Invoice yang sudah ada tidak akan duplikat.')
    add_warning(doc, 'Murid Kids Class Bundle TIDAK mendapat invoice SPP bulanan -- sudah tercover di cicilan 3 termin.')

    add_heading(doc, '3.2  Apply Denda Keterlambatan (Lakukan Tanggal 11+)', level=2)
    add_body(doc, 'Denda Rp 5.000/hari berlaku mulai tanggal 11 untuk invoice yang belum dibayar.')
    add_step(doc, 1, 'Klik menu "Keuangan" > "Daftar Invoice".')
    add_step(doc, 2, 'Klik tombol "Apply Denda".')
    add_step(doc, 3, 'Pilih Tahun dan Bulan.')
    add_step(doc, 4, 'Klik "Apply". Sistem menambah item DENDA ke setiap invoice yang belum lunas.')
    add_note(doc, 'Denda dihitung otomatis: Rp 5.000 x (hari ini - 10). Contoh: bayar tanggal 15, denda Rp 25.000.')

    add_heading(doc, '3.3  Kalkulasi Honor Guru', level=2)
    add_body(doc, 'Honor dihitung berdasarkan sesi yang sudah diisi status absensinya.')
    add_step(doc, 1, 'Klik menu "Honor Guru".')
    add_step(doc, 2, 'Klik tombol "Kalkulasi Honor".')
    add_step(doc, 3, 'Pilih Tahun dan Bulan.')
    add_step(doc, 4, 'Klik "Hitung". Sistem membuat/memperbarui slip honor untuk semua guru.')
    add_note(doc, 'Kalkulasi dilakukan H-2 sebelum akhir bulan (contoh: tanggal 28 untuk bulan 30 hari). Lakukan di hari yang tepat agar semua sesi sudah ter-input.')

    add_heading(doc, '3.4  Cetak & Distribusi Slip Honor Guru', level=2)
    add_step(doc, 1, 'Klik menu "Honor Guru" > pilih guru > klik "Detail Slip".')
    add_step(doc, 2, 'Periksa rincian: honor pokok, honor transport (isi jika ada), honor lain-lain + keterangan.')
    add_step(doc, 3, 'Klik "Cetak Slip". Halaman A4 akan terbuka.')
    add_step(doc, 4, 'Tekan Ctrl+P > simpan sebagai PDF atau cetak langsung.')
    add_step(doc, 5, 'Serahkan slip ke guru atau kirim via WhatsApp.')
    add_warning(doc, 'Hanya Owner yang bisa menandai slip sebagai "Dibayar". Setelah ditandai dibayar, slip terkunci dan tidak bisa diedit.')


def bab4_insidental(doc):
    add_heading(doc, 'Bab 4 - Tugas Insidental', level=1)

    add_heading(doc, '4.1  Daftar Murid Baru (Status Calon)', level=2)
    add_step(doc, 1, 'Klik menu "Murid" > tombol "Tambah Murid".')
    add_step(doc, 2, 'Isi semua data yang wajib: Nama Lengkap, Gender, Tanggal Lahir.')
    add_step(doc, 3, 'Isi data orang tua/wali: Nama, No. HP, Hubungan.')
    add_step(doc, 4, 'Email murid bersifat opsional (boleh kosong).')
    add_step(doc, 5, 'Klik "Simpan". Murid terdaftar dengan status Calon dan kode otomatis (M-YYYY-NNNN).')

    add_heading(doc, '4.2  Jadwalkan Sesi Trial', level=2)
    add_step(doc, 1, 'Buka detail murid (status Calon).')
    add_step(doc, 2, 'Klik tombol "Jadwalkan Trial".')
    add_step(doc, 3, 'Isi: Tanggal Trial, Jam, Guru, Paket yang diminati.')
    add_step(doc, 4, 'Klik "Simpan". Status murid berubah ke Trial. Sesi trial terbuat.')
    add_note(doc, 'Semua sesi trial berdurasi 30 menit, apapun paket yang diminati.')

    add_heading(doc, '4.3  Konversi Trial ke Aktif', level=2)
    add_step(doc, 1, 'Buka detail murid (status Trial).')
    add_step(doc, 2, 'Klik tombol "Konversi ke Aktif".')
    add_step(doc, 3, 'Pilih: Paket, Guru, Ruangan, Hari, Jam jadwal tetap.')
    add_step(doc, 4, 'Klik "Konfirmasi". Status murid berubah ke Aktif.')
    add_step(doc, 5, 'Sistem otomatis membuat Invoice Registrasi Rp 250.000 + Invoice SPP bulan berjalan.')
    add_step(doc, 6, 'Minta murid/orang tua segera melunasi kedua invoice tersebut.')

    add_heading(doc, '4.4  Skip Trial -- Langsung Aktif', level=2)
    add_step(doc, 1, 'Buka detail murid (status Calon).')
    add_step(doc, 2, 'Klik tombol "Skip Trial & Aktifkan".')
    add_step(doc, 3, 'Pilih alasan: Walk-in (datang langsung), Migrasi, Reaktivasi, atau Lulus Kids Class.')
    add_step(doc, 4, 'Isi data kelas: Paket, Guru, Ruangan, Hari, Jam.')
    add_step(doc, 5, 'Klik "Konfirmasi". Murid langsung Aktif + invoice terbuat.')

    add_heading(doc, '4.5  Pengajuan Cuti Murid', level=2)
    add_step(doc, 1, 'Buka detail murid (status Aktif).')
    add_step(doc, 2, 'Klik tombol "Ajukan Cuti".')
    add_step(doc, 3, 'Isi tanggal mulai dan tanggal berakhir cuti (maks 1 bulan).')
    add_step(doc, 4, 'Klik "Simpan". Status murid berubah ke Cuti.')
    add_step(doc, 5, 'Tambahkan biaya cuti Rp 100.000 ke tagihan murid via "Tambah Item" di invoice.')
    add_warning(doc, 'Cuti hanya bisa diperpanjang 1 kali (total maksimal 2 bulan). Murid cuti tidak mendapat tagihan SPP tapi tetap dikenai biaya cuti.')

    add_heading(doc, '4.6  Reschedule Sesi (Izin Murid)', level=2)
    add_body(doc, 'Reschedule hanya diperbolehkan jika: (1) murid memberi info minimal 5 jam sebelum sesi, DAN (2) ini adalah izin pertama murid di bulan tersebut.')
    add_step(doc, 1, 'Input absensi sesi > pilih status IZIN_RESCHEDULE > simpan.')
    add_step(doc, 2, 'Tombol "Buat Sesi Pengganti" akan aktif.')
    add_step(doc, 3, 'Klik tombol tersebut. Modal kecil terbuka.')
    add_step(doc, 4, 'Isi: Tanggal Pengganti, Jam, Ruangan.')
    add_step(doc, 5, 'Klik "Simpan". Sistem cek konflik guru dan ruangan secara otomatis.')
    add_note(doc, 'Izin ke-2 dan seterusnya di bulan yang sama tidak bisa reschedule -- gunakan status IZIN_VIDEO (sesi dianggap hadir, guru tetap dapat honor).')

    add_heading(doc, '4.7  Tambah Kelas Baru untuk Murid Existing (Multi-Kelas)', level=2)
    add_step(doc, 1, 'Buka detail murid (status Aktif).')
    add_step(doc, 2, 'Klik tab "Kelas".')
    add_step(doc, 3, 'Klik tombol "Tambah Kelas".')
    add_step(doc, 4, 'Pilih: Paket baru, Guru, Ruangan, Hari, Jam.')
    add_step(doc, 5, 'Klik "Simpan". Enrollment baru terbuat dengan status ACTIVE.')
    add_note(doc, 'Invoice SPP otomatis hanya untuk kelas utama (primary enrollment). Invoice kelas tambahan dibuat manual jika diperlukan.')

    add_heading(doc, '4.8  Hentikan Kelas / Murid Mundur', level=2)
    add_body(doc, 'Untuk hentikan salah satu kelas (enrollment non-primary):')
    add_step(doc, 1, 'Tab "Kelas" di detail murid > klik "Hentikan" pada enrollment yang dituju.')
    add_step(doc, 2, 'Isi alasan. Klik "Konfirmasi". Enrollment berubah COMPLETED.')
    add_body(doc, 'Untuk hentikan semua kelas (murid mundur):')
    add_step(doc, 3, 'Buka detail murid > klik tombol "Mundurkan Murid".')
    add_step(doc, 4, 'Isi alasan pengunduran diri. Klik "Konfirmasi".')
    add_step(doc, 5, 'Status murid berubah ke Mengundurkan Diri. Semua enrollment COMPLETED.')
    add_warning(doc, 'Jika murid mundur dan ingin kembali lagi di kemudian hari, mereka WAJIB membayar ulang biaya registrasi Rp 250.000.')

    add_heading(doc, '4.9  Generate Cicilan Kids Class Bundle', level=2)
    add_body(doc, 'Khusus untuk murid Kids Class yang memilih opsi bayar cicilan 3 termin.')
    add_step(doc, 1, 'Buka detail murid (Aktif, Kids Class Bundle).')
    add_step(doc, 2, 'Klik tombol "Generate Cicilan Bundle".')
    add_step(doc, 3, 'Isi Tanggal Mulai Program (format YYYY-MM-DD, contoh: 2026-06-01).')
    add_step(doc, 4, 'Klik "Konfirmasi". Tiga invoice cicilan terbuat sekaligus.')
    add_note(doc, 'Cicilan terbagi: Termin 1 (bulan ke-1), Termin 2 (bulan ke-2), Termin 3 (bulan ke-4). Total ketiga termin = Rp 2.180.000.')
    add_warning(doc, 'Generate cicilan hanya bisa dilakukan SEKALI per murid. Sistem akan menolak jika cicilan sudah pernah dibuat.')


def bab5_referensi(doc):
    add_heading(doc, 'Bab 5 - Referensi Cepat', level=1)

    add_heading(doc, '5.1  Status Murid dan Transisi yang Valid', level=2)
    headers = ['Status', 'Arti', 'Transisi yang Mungkin']
    rows = [
        ('Calon',               'Sudah daftar, belum trial',        'ke Trial (jadwal trial) | ke Aktif (skip trial)'),
        ('Trial',               'Sedang/sudah trial, belum aktif',  'ke Aktif (konversi) | ke Mengundurkan Diri'),
        ('Aktif',               'Murid berjalan, kena tagihan SPP', 'ke Cuti | ke Mengundurkan Diri | ke Selesai'),
        ('Cuti',                'Sedang cuti berbayar, sesi pause', 'ke Aktif (cuti berakhir) | ke Mengundurkan Diri'),
        ('Selesai',             'Lulus Kids Class 6 bulan',         'ke Aktif (re-enroll privat, tanpa biaya reg)'),
        ('Mengundurkan Diri',   'Keluar dari studio',               'ke Aktif (re-aktivasi, bayar reg ulang Rp 250.000)'),
    ]
    add_simple_table(doc, headers, rows, [3.0, 4.0, 7.0])

    add_heading(doc, '5.2  Format Nomor Dokumen', level=2)
    headers = ['Jenis Dokumen', 'Format', 'Contoh']
    rows = [
        ('Invoice',         'INV/YYYY/MM/NNNN',  'INV/2026/05/0001'),
        ('Kuitansi',        'KW/YYYY/MM/NNNN',   'KW/2026/05/0001'),
        ('Slip Honor Guru', 'SLIP/YYYY/MM/NNNN', 'SLIP/2026/05/0001'),
        ('Kode Murid',      'M-YYYY-NNNN',       'M-2026-0001'),
    ]
    add_simple_table(doc, headers, rows, [4.5, 4.5, 4.5])

    add_heading(doc, '5.3  Kode Honor Guru', level=2)
    headers = ['Kode', 'Skenario', 'Nominal']
    rows = [
        ('H_REG',    'Sesi reguler (hadir/telat/no-show/hangus)',           'harga paket x 50% / 4'),
        ('H_TRIAL',  'Sesi trial -- murid hadir',                           'Sama seperti H_REG'),
        ('TRIAL_NS', 'Sesi trial -- murid no-show',                         'Rp 0'),
        ('H_VIDEO',  'Izin video pengganti (izin ke-2+)',                   'Sama seperti H_REG'),
        ('H_LIBUR',  'Sesi libur nasional (guru tetap dibayar)',            'Sama seperti H_REG'),
        ('H_HANGUS', 'Murid tidak hadir tanpa konfirmasi',                  'Sama seperti H_REG (guru tetap dibayar)'),
        ('H_PENG',   'Sesi diajar guru pengganti (honor ke pengganti)',     'Sama seperti H_REG'),
        ('H_KIDS',   'Sesi Kids Class (per murid dalam grup)',              'Rp 42.500 x jumlah murid'),
        ('H_UJIAN',  'Pengawas ujian grade',                               'Rp 250.000 flat'),
        ('H_IZIN',   'Sesi original IZIN_RESCHEDULE',                      'Rp 0 (dibayar via sesi pengganti)'),
    ]
    add_simple_table(doc, headers, rows, [2.5, 7.0, 4.0])

    add_heading(doc, '5.4  Cara Eskalasi Masalah ke Owner', level=2)
    add_body(doc, 'Jika menemukan masalah yang tidak bisa diselesaikan sendiri, hubungi owner via WhatsApp dengan format:')
    add_body(doc, '"[SISTEM] [Modul] - Masalah singkat - Langkah yang sudah dicoba"', indent=True)
    add_body(doc, 'Contoh:')
    add_body(doc, '"[SISTEM] M05 - Invoice SPP bulan Juni tidak muncul untuk murid Ahmad - sudah coba generate ulang, tetap tidak ada"', indent=True)
    add_body(doc, '')
    add_body(doc, 'Tangkap screenshot error jika ada (Windows+Shift+S atau PrtScn), kirim bersama pesan WhatsApp.')

    add_heading(doc, '5.5  Komponen Tagihan Manual (dari Katalog)', level=2)
    headers = ['Kode', 'Nama', 'Nominal']
    rows = [
        ('REG',     'Biaya Registrasi',        'Rp 250.000'),
        ('SPP',     'SPP Bulanan',              'Sesuai paket'),
        ('CUTI',    'Biaya Cuti',               'Rp 100.000/pengajuan'),
        ('UJI',     'Ujian + Mini Concert',     'Rp 395.000'),
        ('MC',      'Mini Concert saja',        'Rp 295.000'),
        ('KIDS_FP', 'Final Project Kids Class', 'Rp 140.000/murid'),
        ('DENDA',   'Denda Keterlambatan',      'Rp 5.000/hari (mulai tgl 11)'),
        ('DISKON',  'Diskon Manual',            'Nominal Rp atau % dari item'),
    ]
    add_simple_table(doc, headers, rows, [2.0, 5.0, 4.5])
    add_note(doc, 'Tambah item manual ke invoice via halaman detail invoice > tombol "Tambah Item". Pilih dari daftar katalog di atas. Untuk item DISKON, wajib isi alasan diskon.')


# ─── Main ──────────────────────────────────────────────────────────────────────

def build_document():
    doc = Document()

    # Margin
    for section in doc.sections:
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.0)

    # Cover
    doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run('Panduan Penggunaan Sistem\nRole: Admin')
    set_font(r, size=20, bold=True, color=(0x1F, 0x49, 0x7D))

    doc.add_paragraph()
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = sub.add_run('Musik KITA')
    set_font(r2, size=14, bold=True)

    doc.add_paragraph()
    meta_items = [
        ('Versi',         'v1.0'),
        ('Tanggal',       date.today().strftime('%d %B %Y')),
        ('Berlaku untuk', 'Role: Admin'),
        ('Catatan',       'Kids Class Bundle & Monthly dikelola manual sementara'),
    ]
    for label, value in meta_items:
        p = doc.add_paragraph()
        r1 = p.add_run(f'{label:<18}: ')
        set_font(r1, size=11, bold=True)
        r2 = p.add_run(value)
        set_font(r2, size=11)

    doc.add_page_break()

    # Daftar Isi
    add_heading(doc, 'Daftar Isi', level=1)
    toc_items = [
        ('Bab 1 - Memulai',           '1.1 Login  |  1.2 Tampilan Utama  |  1.3 Owner vs Admin'),
        ('Bab 2 - Tugas Harian',      '2.1 Input Absensi  |  2.2 Catat Pembayaran  |  2.3 Cek Tagihan'),
        ('Bab 3 - Tugas Awal Bulan',  '3.1 Generate SPP  |  3.2 Apply Denda  |  3.3 Kalkulasi Honor  |  3.4 Cetak Slip'),
        ('Bab 4 - Tugas Insidental',  '4.1-4.9: Daftar murid, trial, konversi, cuti, reschedule, multi-kelas, mundur, bundle'),
        ('Bab 5 - Referensi Cepat',   'Status murid, format nomor dokumen, kode honor, eskalasi, katalog tagihan'),
    ]
    for chapter, detail in toc_items:
        p = doc.add_paragraph()
        r1 = p.add_run(chapter)
        set_font(r1, bold=True)
        r2 = p.add_run(f'\n    {detail}')
        set_font(r2, size=9.5, italic=True, color=(0x60, 0x60, 0x60))

    doc.add_page_break()

    # Bab-bab
    bab1_memulai(doc)
    doc.add_page_break()
    bab2_harian(doc)
    doc.add_page_break()
    bab3_awal_bulan(doc)
    doc.add_page_break()
    bab4_insidental(doc)
    doc.add_page_break()
    bab5_referensi(doc)

    return doc


if __name__ == '__main__':
    doc = build_document()
    out = os.path.abspath(OUTPUT_PATH)
    doc.save(out)
    print(f'User Manual Admin berhasil dibuat: {out}')
    print(f'   5 bab: Memulai, Tugas Harian, Awal Bulan, Insidental, Referensi Cepat')
